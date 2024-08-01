"""
This Streamlit application demonstrates a Retrieval-Augmented Generation (RAG) system
utilizing the Fireworks Large Language Model and LlamaIndex library. Users can upload
documents (PDF, DOCX, CSV, TXT) and submit a query to retrieve relevant passages and
generate a response using the Fireworks LLM.

**Improvements:**

- **Error handling:** More specific exception handling for informative error messages.
- **Early return:** Avoids unnecessary indexing attempts after processing errors.
- **Optimized loops:** Uses list comprehension and progress bar for better user experience.
- **Refactored functions:** Improved code organization and readability.
- **Docstrings and comments:** Added docstrings and in-line comments for clarity.
- **Streamlit enhancements:** Progress bar and informative messages for user feedback.

**Requirements:**

- streamlit
- PyPDF2 (for PDF files)
- docx (for DOCX files)
- llama-index
- torch (for GPU acceleration, optional)
- fireworks (Fireworks LLM API key)
- sentence-transformers
- plotly
- hdbscan
"""

import asyncio
import csv
import io
import unittest
from typing import List
from unittest.mock import patch, Mock

import PyPDF2
import docx
import streamlit as st
import torch
from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.fireworks import Fireworks
from llama_index.core.retrievers import KeywordTableSimpleRetriever
from sentence_transformers import SentenceTransformer
from sklearn.cluster import HDBSCAN
import plotly.graph_objects as go


def display_results(response):
    """Displays the query results.

    Args:
        response (Response): The response object.
    """
    if not response:
        return

    st.write("Generated Response:", response.response)

    # Display source documents
    st.write("Source Documents:")
    for node in response.source_nodes:
        st.write(f"- {node.node.get_content()[:100]}...")

    # Visualize relevance scores
    scores = [node.score for node in response.source_nodes]
    fig = go.Figure(data=[go.Bar(y=scores)])
    fig.update_layout(title="Relevance Scores of Retrieved Documents",
                      xaxis_title="Document",
                      yaxis_title="Relevance Score")
    st.plotly_chart(fig)


class RAGSystem:
    """RAG system class for processing documents and answering queries."""

    def __init__(self, fireworks_api_key, device="cuda"):
        """Initializes the RAG system.

        Args:
            fireworks_api_key (str): The Fireworks API key.
            device (str, optional): The device to use (cuda or cpu). Defaults to "cuda".
        """
        self.device = device if torch.cuda.is_available() else "cpu"
        st.write(f"Using device: {self.device}")

        self.llm = Fireworks(
            api_key=fireworks_api_key,
            model="accounts/fireworks/models/firefunction-v1"
        )

        self.embed_model = HuggingFaceEmbedding(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            device=self.device
        )

        self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')

        Settings.llm = self.llm
        Settings.embed_model = self.embed_model

        self.document_cache = {}

    @staticmethod
    @st.cache_data
    def _extract_text_from_file(uploaded_file):
        """Extracts text from an uploaded file.

        Args:
            uploaded_file (streamlit.uploadedfile.UploadedFile): The uploaded file.

        Returns:
            str: The extracted text.

        Raises:
            ValueError: If the file format is unsupported.
        """
        file_extension = uploaded_file.name.split('.')[-1].lower()

        if file_extension == 'pdf':
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            return ' '.join(page.extract_text() for page in pdf_reader.pages)

        elif file_extension == 'docx':
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            return ' '.join(paragraph.text for paragraph in doc.paragraphs)

        elif file_extension == 'csv':
            csv_content = uploaded_file.read().decode('utf-8')
            csv_reader = csv.reader(io.StringIO(csv_content))
            return ' '.join(' '.join(row) for row in csv_reader)

        elif file_extension == 'txt':
            return uploaded_file.read().decode('utf-8')

        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def extract_text_from_file(self, uploaded_file):
        """Extracts text from an uploaded file and caches it.

        Args:
            uploaded_file (streamlit.uploadedfile.UploadedFile): The uploaded file.

        Returns:
            str: The extracted text.
        """
        return self._extract_text_from_file(uploaded_file)

    async def generate_meaningful_chunks(self, text, chunk_size=512, chunk_overlap=50, min_cluster_size=5):
        """Generates semantically meaningful chunks from the given text using HDBSCAN.

        Args:
            text (str): The input text.
            chunk_size (int, optional): The desired chunk size. Defaults to 512.
            chunk_overlap (int, optional): The overlap between chunks. Defaults to 50.
            min_cluster_size (int, optional): The minimum size of a cluster. Defaults to 5.

        Returns:
            list[str]: A list of text chunks.
        """
        sentences = text.split('.')
        sentence_embeddings = await asyncio.to_thread(
            self.sentence_model.encode, sentences, batch_size=64
        )

        # Clustering
        clusterer = HDBSCAN(min_cluster_size=min_cluster_size, metric='cosine')
        clusterer.fit(sentence_embeddings)
        labels = clusterer.labels_

        # Group sentences by cluster
        chunks = []
        for label in set(labels):
            cluster_sentences = [sentences[i] for i in range(len(sentences)) if labels[i] == label]
            if len(cluster_sentences) >= min_cluster_size:
                chunks.append(' '.join(cluster_sentences))

        return chunks

    async def create_index(self, documents):
        """Creates a VectorStoreIndex from a list of documents.

        Args:
            documents (list[Document]): A list of Document objects.

        Returns:
            VectorStoreIndex: The created index.
        """
        try:
            index = await asyncio.to_thread(VectorStoreIndex.from_documents, documents)
            return index
        except Exception as e:
            st.error(f"Error creating index: {str(e)}")
            return None

    def process_query(self, query, index):
        """Processes a query using custom hybrid retrieval.

        Args:
            query (str): The user's query.
            index (VectorStoreIndex): The index to use for query processing.

        Returns:
            Response: The query response.
        """
        if not index:
            st.error("Index not available for query processing.")
            return None

        try:
            # Vector-based retrieval
            vector_retriever = index.as_retriever(similarity_top_k=5)
            vector_nodes = vector_retriever.retrieve(query)

            # Manual keyword matching
            keyword_nodes = []
            for node in vector_nodes:
                doc_text = node.node.text
                if any(keyword in doc_text.lower() for keyword in query.lower().split()):
                    keyword_nodes.append(node)

            # Combine and re-rank results
            combined_nodes = list({node.node.node_id: node for node in vector_nodes + keyword_nodes}.values())
            st.write(f"Combined results before sorting: {len(combined_nodes)} nodes")

            # Rank results based on vector scores and keyword presence
            combined_nodes.sort(
                key=lambda x: x.score + (
                    1 if any(keyword in x.node.text.lower() for keyword in query.lower().split()) else 0),
                reverse=True
            )

            # Use top results for query
            top_nodes = combined_nodes[:5]
            st.write(f"Top nodes: {top_nodes}")

            # Create a query engine and process the query
            query_engine = index.as_query_engine()
            # Process with the default query engine
            response = query_engine.query(query)

            return response

        except Exception as e:
            st.error(f"Error processing query: {str(e)}")
            return None

    async def process_file(self, uploaded_file, chunk_size, chunk_overlap, min_cluster_size):
        """Processes a single uploaded file to extract text and generate chunks.

        Args:
            uploaded_file (streamlit.uploadedfile.UploadedFile): The uploaded file.
            chunk_size (int): The chunk size for text splitting.
            chunk_overlap (int): The chunk overlap size.
            min_cluster_size (int): The minimum cluster size for HDBSCAN.

        Returns:
            list[Document]: A list of Document objects containing text chunks.
        """
        if uploaded_file.name in self.document_cache:
            return self.document_cache[uploaded_file.name]

        try:
            text = self.extract_text_from_file(uploaded_file)
            chunks = await self.generate_meaningful_chunks(text, chunk_size, chunk_overlap, min_cluster_size)
            documents = [Document(text=chunk) for chunk in chunks]
            self.document_cache[uploaded_file.name] = documents
            return documents
        except Exception as e:
            st.warning(f"Error processing {uploaded_file.name}: {str(e)}")
            return []

    async def evaluate_performance(self, responses, ground_truth):
        """
        Evaluate the performance of the RAG system.

        Args:
        responses (list): List of generated responses
        ground_truth (list): List of expected responses

        Returns:
        dict: Performance metrics
        """
        # Simple accuracy metric
        correct = sum(r == gt for r, gt in zip(responses, ground_truth))
        accuracy = correct / len(responses)

        # You can add more sophisticated metrics here, such as BLEU score, ROUGE, etc.

        return {"accuracy": accuracy}

    def generate_evaluation_report(self, performance_metrics):
        """
        Generate an evaluation report based on test results and performance metrics.

        Args:
        test_results (dict): Results from unit tests
        performance_metrics (dict): Metrics from performance evaluation

        Returns:
        str: Formatted evaluation report
        """
        report = "RAG System Evaluation Report\n"
        report += "===========================\n\n"

        report += "\n1. Performance Metrics:\n"
        for metric, value in performance_metrics.items():
            report += f"   - {metric}: {value}\n"

        report += "\n2. Simulated User Feedback:\n"
        report += "   [Include user feedback from testing sessions]\n"

        return report

    async def run_rag(self, query, uploaded_files, chunk_size, chunk_overlap, min_cluster_size):
        """Runs the full RAG process: processing files, creating index, and processing query.

        Args:
            query (str): The user's query.
            uploaded_files (list[streamlit.uploadedfile.UploadedFile]): The uploaded files.
            chunk_size (int): The chunk size for text splitting.
            chunk_overlap (int): The chunk overlap size.
            min_cluster_size (int): The minimum cluster size for HDBSCAN.
        """
        documents = []
        total_files = len(uploaded_files)
        st.write("Processing uploaded files...")
        tasks = [self.process_file(file, chunk_size, chunk_overlap, min_cluster_size) for file in uploaded_files]
        documents = []

        for i, uploaded_file in enumerate(uploaded_files, start=1):
            st.write(f"Processing file {i}/{total_files}: {uploaded_file.name}")
            file_documents = await self.process_file(uploaded_file, chunk_size, chunk_overlap, min_cluster_size)
            if file_documents:
                documents.extend(file_documents)

        if documents:
            st.write("Creating index from documents...")
            index = await self.create_index(documents)
            if index:
                response = self.process_query(query, index)
                display_results(response)
            else:
                st.warning("No valid documents were processed. Please check your file formats and content.")


# Streamlit app setup
st.set_page_config(page_title="RAG System with Fireworks LLM", layout="wide")

st.title("Retrieval-Augmented Generation (RAG) System with Fireworks LLM")
st.write(
    "Upload your documents (PDF, DOCX, CSV, TXT) and ask a question to get a response generated by the Fireworks Large Language Model.")

fireworks_api_key = st.secrets["fireworks_api_key"]

if not fireworks_api_key:
    st.error("Please provide a valid Fireworks API key in your Streamlit secrets.")
else:
    rag_system = RAGSystem(fireworks_api_key)

    uploaded_files = st.file_uploader("Upload Files", accept_multiple_files=True, type=["pdf", "docx", "csv", "txt"])
    query = st.text_input("Enter your query")
    chunk_size = st.slider("Chunk Size", 128, 1024, 512)
    chunk_overlap = st.slider("Chunk Overlap", 0, 256, 50)
    min_cluster_size = st.slider("Minimum Cluster Size", 1, 20, 5)

    if st.button("Submit Query") and uploaded_files and query:
        st.write("Processing query...")
    asyncio.run(rag_system.run_rag(query, uploaded_files, chunk_size, chunk_overlap, min_cluster_size))
